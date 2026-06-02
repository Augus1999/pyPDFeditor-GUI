# -*- coding: utf-8 -*-
"""Convert Distilmark's Markdown output to other formats.

Supported targets:
  • HTML  — uses the ``markdown`` library; produces a clean styled document.
  • DOCX  — prefers Pandoc when available, falls back to a basic python-docx
            converter that handles headings, lists, paragraphs and images.
  • Combined Markdown — concatenates multiple ``.md`` files with separators.
"""
from __future__ import annotations

import re
import shutil
import subprocess
from pathlib import Path

# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------

_HTML_CSS = """
body{font-family:-apple-system,'Segoe UI Variable','Segoe UI',Inter,sans-serif;
     max-width:860px;margin:32px auto;padding:0 24px;color:#1e2230;
     line-height:1.6;background:#f6f7fb;}
h1,h2,h3,h4,h5,h6{font-weight:700;line-height:1.25;margin-top:1.6em;color:#0f1117;}
h1{font-size:2em;border-bottom:1px solid #d9dce6;padding-bottom:.3em;}
h2{font-size:1.5em;}
code{background:#e6e8f0;padding:2px 6px;border-radius:4px;font-size:0.93em;}
pre{background:#1e2230;color:#e7eaf0;padding:14px 16px;border-radius:10px;
    overflow-x:auto;}
pre code{background:transparent;color:inherit;padding:0;}
blockquote{border-left:3px solid #2563eb;padding:.4em 1em;color:#3b4252;
           background:#eef0f6;margin:1em 0;border-radius:0 8px 8px 0;}
table{border-collapse:collapse;margin:1em 0;}
th,td{border:1px solid #d9dce6;padding:8px 12px;}
th{background:#eef0f6;}
img{max-width:100%;border-radius:8px;}
a{color:#2563eb;}
hr{border:none;border-top:1px solid #d9dce6;margin:2em 0;}
"""


def to_html(md_text: str, out_path: Path, title: str | None = None) -> Path:
    """Convert Markdown text to a styled standalone HTML file."""
    import markdown  # type: ignore
    html_body = markdown.markdown(
        md_text,
        extensions=["extra", "tables", "fenced_code", "toc", "sane_lists"],
        output_format="html5",
    )
    page = (
        "<!doctype html>\n<html><head>"
        "<meta charset='utf-8'>"
        f"<title>{title or out_path.stem}</title>"
        f"<style>{_HTML_CSS}</style>"
        "</head><body>\n"
        f"{html_body}\n"
        "</body></html>"
    )
    out_path.write_text(page, encoding="utf-8")
    return out_path


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _pandoc_available() -> str | None:
    return shutil.which("pandoc")


def to_docx(md_text: str, out_path: Path, md_source: Path | None = None) -> Path:
    """Convert Markdown to DOCX.

    Prefers Pandoc (best quality); falls back to a basic python-docx converter
    that handles headings, paragraphs, lists, blockquotes and images.
    """
    pandoc = _pandoc_available()
    if pandoc:
        # When we have a source .md on disk, run pandoc against it so image
        # paths resolve against the file's directory.
        if md_source is not None and md_source.exists():
            subprocess.run(
                [pandoc, str(md_source), "-o", str(out_path)],
                check=True, capture_output=True,
            )
        else:
            subprocess.run(
                [pandoc, "-f", "markdown", "-o", str(out_path)],
                input=md_text.encode("utf-8"), check=True, capture_output=True,
            )
        return out_path
    return _fallback_to_docx(md_text, out_path, md_source)


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_ULIST_RE = re.compile(r"^[\-*]\s+(.*)$")
_OLIST_RE = re.compile(r"^\d+\.\s+(.*)$")
_QUOTE_RE = re.compile(r"^>\s?(.*)$")
_IMAGE_RE = re.compile(r"^!\[[^\]]*\]\(([^)]+)\)\s*$")


def _fallback_to_docx(md_text: str, out_path: Path, md_source: Path | None) -> Path:
    from docx import Document  # type: ignore
    from docx.shared import Inches  # type: ignore

    base = md_source.parent if md_source else None
    doc = Document()

    in_code = False
    code_buf: list[str] = []
    for raw in md_text.splitlines():
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                code = "\n".join(code_buf)
                p = doc.add_paragraph()
                run = p.add_run(code)
                run.font.name = "Consolas"
                code_buf.clear()
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_buf.append(raw)
            continue
        if not line.strip():
            doc.add_paragraph("")
            continue
        m = _HEADING_RE.match(line)
        if m:
            level = min(len(m.group(1)), 4)
            doc.add_heading(m.group(2).strip(), level=level)
            continue
        m = _IMAGE_RE.match(line)
        if m:
            path = m.group(1)
            full = Path(path) if Path(path).is_absolute() else (base / path if base else Path(path))
            try:
                doc.add_picture(str(full), width=Inches(5.5))
            except Exception:
                doc.add_paragraph(f"[image: {path}]")
            continue
        m = _ULIST_RE.match(line)
        if m:
            doc.add_paragraph(m.group(1).strip(), style="List Bullet")
            continue
        m = _OLIST_RE.match(line)
        if m:
            doc.add_paragraph(m.group(1).strip(), style="List Number")
            continue
        m = _QUOTE_RE.match(line)
        if m:
            p = doc.add_paragraph(m.group(1).strip(), style="Intense Quote")
            continue
        doc.add_paragraph(line)
    doc.save(str(out_path))
    return out_path


# ---------------------------------------------------------------------------
# Combined Markdown
# ---------------------------------------------------------------------------

def combine_markdown(md_paths: list[Path], out_path: Path, header: bool = True) -> Path:
    """Concatenate several .md files into one, with a section per source."""
    parts: list[str] = []
    for p in md_paths:
        try:
            body = p.read_text(encoding="utf-8")
        except OSError:
            continue
        if header:
            parts.append(f"# {p.stem}\n")
        parts.append(body.strip())
        parts.append("\n\n---\n")
    out_path.write_text("\n".join(parts).rstrip() + "\n", encoding="utf-8")
    return out_path
